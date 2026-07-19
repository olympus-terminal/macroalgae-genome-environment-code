#!/usr/bin/env python3
"""Build the GEE-primary Supplemental Information legends package (V7).

This presentation builder uses only authenticated completed analysis outputs. It
does not simulate, impute, or recompute scientific results. It performs four
versioned operations:

1. revise the V6 legend document to place named Google Earth Engine variables
   first and AlphaEarth Foundations analyses second;
2. render Figure S3 from the retained non-null result table and the corrected
   99,999-permutation conditional-tail result table;
3. copy the existing exploratory AEF--Pfam landscape byte-for-byte as Figure S4;
4. render DOCX and Figure S3 supporting-page PDFs and write an integrity record.

Generated: 2026-07-19
Runtime: Python 3.10; python-docx, lxml, pandas, Matplotlib, LibreOffice,
Tectonic, qpdf, Poppler.
"""

from __future__ import annotations

import hashlib
import importlib.util
import json
import os
import platform
import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib
import numpy as np
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from lxml import etree
from openpyxl import load_workbook


STAMP = "20260719_211018"
SCRIPT = Path(__file__).resolve()
ROOT = SCRIPT.parents[2]
SUPPLEMENT = SCRIPT.parent
STATS = ROOT / "ISCIENCE_REVISION_20260711" / "analysis_stats"
GEE_VALIDATION = ROOT / "ISCIENCE_REVISION_20260711" / "gee_validation"
AEF_RUN = (
    ROOT
    / "ISCIENCE_REVISION_20260711"
    / "aef"
    / "full_aef_corrected_run_20260718_154213"
)

SOURCE_STEM = (
    SUPPLEMENT
    / "Supplemental_Information_revised_legends_20260718_234005_V6"
)
SOURCE_DOCX = SOURCE_STEM.with_suffix(".docx")
SOURCE_TXT = SOURCE_STEM.with_suffix(".txt")
SOURCE_PDF = SOURCE_STEM.with_suffix(".pdf")
SOURCE_TABLE_S2 = (
    SUPPLEMENT / "Table_S2_revised_exact_id_GEE_validation_20260719_211551_V3.xlsx"
)
SOURCE_TABLE_S3 = SUPPLEMENT / "Table_S3_AEF_20260719_211551_V6.xlsx"

SOURCE_PLOT_SCRIPT = STATS / "run_robustness_20260711_085930.py"
SOURCE_AEF_NON_NULL = STATS / "archived_AEF_priority_robustness_20260711_131706.csv"
SOURCE_AEF_NULL = STATS / "corrected_AEF_structured_null_conditional_tail_20260719_210254.csv"
SOURCE_AEF_CANDIDATES = STATS / "corrected_AEF_robust_candidate_table_20260719_210254.csv"
SOURCE_TEX_BUILDER = SUPPLEMENT / "build_figure_s3_supporting_package_20260711_175934.py"
SOURCE_FIGURE_S4_PDF = AEF_RUN / "Figure5_corrected_AEF_Pfam_landscape.pdf"
SOURCE_FIGURE_S4_SVG = AEF_RUN / "Figure5_corrected_AEF_Pfam_landscape.svg"

OUTPUT_STEM = (
    SUPPLEMENT
    / f"Supplemental_Information_revised_legends_{STAMP}_V7"
)
OUTPUT_DOCX = OUTPUT_STEM.with_suffix(".docx")
OUTPUT_PDF = OUTPUT_STEM.with_suffix(".pdf")
OUTPUT_TXT = OUTPUT_STEM.with_suffix(".txt")
OUTPUT_AUDIT = SUPPLEMENT / f"Supplemental_Information_revised_legends_{STAMP}_V7_integrity.json"

FIGURE_S3_STEM = f"Figure_S3_posthoc_sensitivity_checks_V3_{STAMP}"
FIGURE_S3_PDF = SUPPLEMENT / f"{FIGURE_S3_STEM}.pdf"
FIGURE_S3_SVG = SUPPLEMENT / f"{FIGURE_S3_STEM}.svg"
FIGURE_S3_CAPTION = SUPPLEMENT / f"Figure_S3_posthoc_sensitivity_checks_caption_V3_{STAMP}.txt"
FIGURE_S3_PAGE_TEX = SUPPLEMENT / f"Figure_S3_posthoc_sensitivity_checks_supporting_page_V4_{STAMP}.tex"
FIGURE_S3_PAGE_PDF = FIGURE_S3_PAGE_TEX.with_suffix(".pdf")

FIGURE_S4_PDF = SUPPLEMENT / f"Figure_S4_exploratory_AEF_Pfam_correlation_profile_landscape_{STAMP}.pdf"
FIGURE_S4_SVG = SUPPLEMENT / f"Figure_S4_exploratory_AEF_Pfam_correlation_profile_landscape_{STAMP}.svg"

SOFFICE_CANDIDATES = (
    Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
    Path("/opt/homebrew/bin/soffice"),
)

MANUSCRIPT_TITLE = (
    "Named Earth-Observation Variables Identify Temperature-Centered "
    "Protein-Family Covariation across Macroalgal Genomes"
)

FIGURE_S3_LEGEND = (
    "Secondary selected-set sensitivity analysis of protein family (Pfam)--AlphaEarth "
    "Foundations (AEF) latent-feature associations. Rank-based effects are shown for "
    "raw counts, counts per total Pfam hit, and counts per final peptide record. The "
    "seven required specifications were unique-site analyses of the three abundance "
    "representations, a total-hit quality/phylum model with coordinate-site-clustered "
    "covariance, a total-hit Benchmarking Universal Single-Copy Orthologs (BUSCO) "
    ">=50% unique-site analysis, a total-hit three-phylogenetic-eigenvector topology "
    "model with site-clustered covariance, and a total-hit structured null. The null "
    "used 99,999 permutations of intact sites within site phylum-composition strata; "
    "its two-sided p value was twice the smaller conditional tail, including ties, and "
    "Benjamini--Hochberg (BH) correction spanned all 68 selected pairs within each "
    "abundance representation. PF01638--A52 and PF01638--A53 retained direction and "
    "selected-family support across all seven specifications (7/7). Their total-hit "
    "site-mean effects were rho = 0.450 and rho = -0.542; corrected structured-null "
    "p/q values were 8.0e-5/0.00272 and 2.0e-5/0.00136, respectively. PF10988--A36 "
    "and PF13411--A18 each passed five of seven specifications (5/7). The displayed "
    "12 pairs are a visualization subset of the 68-pair selected family. A00--A63 "
    "denote unitless latent axes."
)

FIGURE_S4_LEGEND = (
    "Secondary exploratory AEF--Pfam correlation-profile landscape. Average-linkage "
    "hierarchical clustering with correlation distance organizes Pfam--AEF Spearman "
    "profiles across 126 exact-genome-identifier (ID) records. All 10,707 strict Pfam "
    "accessions and 64 axes are displayed as row-wise z scores. The exact-profile "
    "audit identified 8,713 unique profiles and 1,994 additional accessions sharing an "
    "exact profile; all 10,707 accessions remain in the display. The top bar gives the "
    "number of global-BH-supported Pfam pairs per axis, and the right strip gives each "
    "Pfam's maximum absolute Spearman coefficient. AEF axes A00--A63 are unitless "
    "latent features. Figure S3 reports the selected-set AEF sensitivity analysis, and "
    "Table S3G--I reports site-level external alignment with named Google Earth Engine "
    "variables. The alignment is a crosswalk rather than a semantic decoding of the "
    "axes."
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def file_record(path: Path) -> dict[str, Any]:
    return {
        "path": str(path.resolve()),
        "bytes": path.stat().st_size,
        "sha256": sha256(path),
    }


def run_checked(command: list[str], *, cwd: Path | None = None) -> subprocess.CompletedProcess[str]:
    result = subprocess.run(command, cwd=cwd, capture_output=True, text=True, check=False)
    if result.returncode != 0:
        raise RuntimeError(
            f"Command failed ({result.returncode}): {' '.join(command)}\n"
            f"STDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
        )
    return result


def pdf_pages(path: Path) -> int:
    info = run_checked(["pdfinfo", str(path)]).stdout
    match = re.search(r"^Pages:\s+(\d+)$", info, flags=re.MULTILINE)
    if match is None:
        raise RuntimeError(f"Could not determine PDF page count: {path}")
    return int(match.group(1))


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load module: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def display_text(text: str) -> str:
    """Apply publication typography to prose without changing its values."""
    return text.replace("--", "–").replace(">=", "≥")


def replace_paragraph(paragraph: Paragraph, label: str | None, body: str, *, bold_all: bool = False) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag.endswith("}pPr"):
            continue
        element.remove(child)
    if label is not None:
        lead = paragraph.add_run(display_text(label) + " ")
        lead.bold = True
    run = paragraph.add_run(display_text(body))
    if bold_all:
        run.bold = True


def insert_after(anchor: Paragraph, label: str, body: str) -> Paragraph:
    paragraph_xml = OxmlElement("w:p")
    anchor._p.addnext(paragraph_xml)
    paragraph = Paragraph(paragraph_xml, anchor._parent)
    paragraph.style = "Normal"
    replace_paragraph(paragraph, label, body)
    return paragraph


def find_unique(document: Document, prefix: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {prefix!r}; found {len(matches)}")
    return matches[0]


def document_text(document: Document) -> str:
    return "\n\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text) + "\n"


def build_legend_document() -> dict[str, Any]:
    document = Document(SOURCE_DOCX)
    source_paragraphs = len(document.paragraphs)

    replace_paragraph(document.paragraphs[1], None, MANUSCRIPT_TITLE, bold_all=True)

    figure_s3 = find_unique(document, "Figure S3.")
    replace_paragraph(figure_s3, "Figure S3.", FIGURE_S3_LEGEND)
    insert_after(figure_s3, "Figure S4.", FIGURE_S4_LEGEND)

    replace_paragraph(
        find_unique(document, "Table S1G."),
        "Table S1G. Raw Pfam input index.",
        "Path and hash for the 126-genome matrix reconstructed from per-genome profile hidden Markov model "
        "searches performed with HMMER hmmsearch. Raw domain-hit counts are the primary abundance measure. "
        "The Figure 3 recorded-metadata screen comprised 42,828 tests; 491 had p < 0.001, 228 had global "
        "Benjamini–Hochberg (BH) q < 0.05, and 13 met Bonferroni p < 0.05.",
    )

    replace_paragraph(
        find_unique(document, "Table S2."),
        "Table S2. Primary named-variable Google Earth Engine (GEE) discovery and sensitivity analyses.",
        "Exact-genome-identifier (ID) matching joined environmental values and Pfam counts for 126 genomes. "
        "The discovery family comprised 87,123 tests, with 6,700–6,703 estimable Pfams per named variable; "
        "84 associations representing 51 Pfams met global Benjamini--Hochberg (BH) q < 0.05, and 13 met "
        "Bonferroni p < 0.05. Of the 84 selected associations, 49 retained the discovery direction and "
        "selected-family support across all seven required sensitivity specifications; all 13 discovery-"
        "Bonferroni associations passed all seven.",
    )
    replace_paragraph(
        find_unique(document, "Table S2C."),
        "Table S2C. Global-BH discovery associations.",
        "All 84 raw-count Pfam--GEE pairs meeting global BH q < 0.05 in the 87,123-test exact-ID family. "
        "These pairs define the selected family evaluated in Table S2E--G.",
    )
    s2d = find_unique(document, "Table S2D.")
    replace_paragraph(
        s2d,
        "Table S2D. Analysis output index.",
        "Paths, byte sizes, row counts, and 256-bit Secure Hash Algorithm (SHA-256) hashes for the exact-ID "
        "extraction, complete discovery family, selected-pair sensitivity analyses, and structured null.",
    )
    anchor = s2d
    for label, body in [
        (
            "Table S2E. GEE selected-pair sensitivity results.",
            "The 2,016 rows comprise 84 discovery pairs across three abundance representations and eight "
            "specifications: descriptive genome-level Spearman correlation; unique-site Spearman correlation; "
            "a quality/phylum model with site-clustered covariance; BUSCO >=50% and >=70% unique-site analyses; "
            "and three-, five-, and ten-phylogenetic-eigenvector topology models with site-clustered covariance. "
            "BH correction is applied within each specification-by-representation family of 84 pairs.",
        ),
        (
            "Table S2F. GEE structured-null results.",
            "For each of the 84 selected pairs, total-hit-normalized site means were evaluated with 99,999 "
            "permutations of intact site environmental labels within site phylum-composition strata. Two-sided "
            "p values equal twice the smaller conditional randomization tail, including ties, and BH correction "
            "spans all 84 selected pairs. Fifty-six associations met corrected q < 0.05.",
        ),
        (
            "Table S2G. GEE seven-specification candidate summary.",
            "The required specifications are unique-site raw counts, total-hit normalization, peptide "
            "normalization, the total-hit quality/phylum site-cluster model, the total-hit BUSCO >=50% "
            "unique-site analysis, the total-hit three-phylogenetic-eigenvector topology model, and the total-hit "
            "structured null. Forty-nine of 84 selected pairs retained direction and selected-family support in "
            "all seven; all 13 discovery-Bonferroni pairs passed all seven.",
        ),
    ]:
        anchor = insert_after(anchor, label, body)

    replace_paragraph(
        find_unique(document, "Table S3."),
        "Table S3. Secondary AlphaEarth Foundations (AEF) association and site-alignment analyses.",
        "The pooled Pfam--AEF screen evaluated 10,707 strict Pfam accessions against 64 unitless latent axes "
        "across 126 exact-ID genomes (685,248 tests). The site-level external alignment evaluated 64 axes "
        "against 13 named GEE variables at up to 90 coordinate sites (832 tests); 224 met global BH q < 0.05 "
        "and 44 met Bonferroni p < 0.05. AEF axes remain unitless latent features, and Figure S4 displays the "
        "complete exploratory Pfam--AEF correlation-profile landscape.",
    )
    replace_paragraph(
        find_unique(document, "Table S3A."),
        "Table S3A. Pooled genome-level Pfam--AEF associations.",
        "All 24 raw-count Pfam--AEF pairs meeting global BH q < 0.05 in the 685,248-test family, spanning "
        "17 Pfams and 13 AEF axes; three pairs also met Bonferroni p < 0.05.",
    )
    replace_paragraph(
        find_unique(document, "Table S3B."),
        "Table S3B. Phylum-centered genome-level Pfam--AEF associations.",
        "All 17 pairs meeting global BH q < 0.05 after subtracting each phylum-specific Pfam mean before "
        "computing the 126-genome correlations.",
    )
    replace_paragraph(
        find_unique(document, "Table S3C."),
        "Table S3C. Site-level AEF associations with recorded metadata.",
        "All 192 Spearman correlations of A00--A63 with recorded temperature, latitude, and longitude after "
        "giving each of 90 unique coordinate sites one observation. BH and Bonferroni correction span the "
        "192-test family; 51 correlations met BH q < 0.05 and 23 met Bonferroni p < 0.05.",
    )
    replace_paragraph(
        find_unique(document, "Table S3D."),
        "Table S3D. Site-level recorded-metadata summary by AEF axis.",
        "For each unitless axis, the table reports site-level correlations with recorded temperature, latitude, "
        "and longitude, the strongest recorded descriptor, and the pooled Pfam-association summary.",
    )
    replace_paragraph(
        find_unique(document, "Table S3E."),
        "Table S3E. AEF output index.",
        "Paths, byte sizes, row counts, and SHA-256 hashes for the pooled AEF analyses, site-level AEF--GEE "
        "alignment, and corrected selected-set sensitivity outputs.",
    )
    s3f = find_unique(document, "Table S3F.")
    replace_paragraph(
        s3f,
        "Table S3F. Within-phylum specification summary.",
        "Numbers of genomes, eligible Pfams, tests, nominal p < 0.001 results, global BH q < 0.05 results, "
        "and minimum q for the Chlorophyta, Ochrophyta, and Rhodophyta screens and the Rhodophyta >=10-"
        "nonzero sensitivity. The global-BH count was zero in all four specifications.",
    )
    anchor = s3f
    for label, body in [
        (
            "Table S3G. Site-level AEF--GEE alignment.",
            "All 832 Spearman correlations between 64 unitless AEF axes and 13 named GEE variables after exact-"
            "ID joining and unique-site aggregation. Global BH correction spans all 832 correlations; 224 met "
            "q < 0.05 and 44 met Bonferroni p < 0.05. The table is an external crosswalk between latent axes and "
            "named variables rather than a semantic decoder assigning unique physical labels.",
        ),
        (
            "Table S3H. GEE-variable alignment summary.",
            "For each of the 13 named GEE variables, the table reports the number of coordinate sites, the "
            "strongest AEF-axis correlation, and counts meeting global BH and Bonferroni thresholds in the "
            "832-test alignment family.",
        ),
        (
            "Table S3I. AEF-axis alignment summary.",
            "For each of A00--A63, the table reports its strongest named GEE correlate and counts meeting global "
            "BH and Bonferroni thresholds. A52 aligned negatively with all five sea-surface-temperature (SST) "
            "summaries (strongest rho = -0.704 for maximum SST), whereas A53 aligned positively with all five "
            "(rho = 0.604 for mean SST).",
        ),
        (
            "Table S3J. Corrected selected-AEF structured null.",
            "The 204 rows comprise 68 selected Pfam--AEF pairs across raw-count, total-hit-normalized, and "
            "peptide-normalized representations. Each representation used 99,999 intact-site permutations "
            "within site phylum-composition strata. Two-sided p values equal twice the smaller conditional tail, "
            "including ties, and BH correction spans 68 pairs separately within each representation.",
        ),
        (
            "Table S3K. Corrected selected-AEF seven-specification summary.",
            "PF01638--A52 and PF01638--A53 retained direction and selected-family support in all seven required "
            "specifications (7/7). PF10988--A36 and PF13411--A18 each passed five of seven specifications (5/7).",
        ),
    ]:
        anchor = insert_after(anchor, label, body)

    output_text = document_text(document)
    prohibited = [
        r"Figure 5",
        r"(?<!9)9,999 permutations",
        r"PF10988–A36 met six",
        r"empirical p = 0\.0001",
        r"selected-set q = 0\.0034",
        r"separate from the seven-check post hoc AEF sensitivity set",
    ]
    found = [pattern for pattern in prohibited if re.search(pattern, output_text)]
    if found:
        raise RuntimeError(f"Stale legend wording remains: {found}")
    required = [
        MANUSCRIPT_TITLE,
        "Figure S4.",
        "49 retained the discovery direction",
        "Fifty-six associations met corrected q < 0.05",
        "224 met global BH q < 0.05",
        "PF10988–A36 and PF13411–A18 each passed five of seven",
        "semantic decoder",
    ]
    missing = [phrase for phrase in required if phrase not in output_text]
    if missing:
        raise RuntimeError(f"Required V7 wording is absent: {missing}")

    document.save(OUTPUT_DOCX)
    OUTPUT_TXT.write_text(output_text, encoding="utf-8")
    reopened = Document(OUTPUT_DOCX)
    if document_text(reopened) != output_text:
        raise RuntimeError("V7 DOCX visible paragraphs and TXT differ")
    with zipfile.ZipFile(OUTPUT_DOCX) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("V7 DOCX ZIP integrity test failed")
        document_xml = archive.read("word/document.xml")
    if b"<w:ins" in document_xml or b"<w:del" in document_xml:
        raise RuntimeError("Tracked-change markup is present in V7")
    return {
        "source_paragraphs": source_paragraphs,
        "output_paragraphs": len(reopened.paragraphs),
        "inserted_paragraphs": len(reopened.paragraphs) - source_paragraphs,
        "docx_txt_visible_text_identity": True,
        "tracked_change_markup_absent": True,
    }


def build_figure_s3() -> dict[str, Any]:
    source = SOURCE_PLOT_SCRIPT.read_text(encoding="utf-8")
    substitutions = {
        "A  Priority archived latent-feature associations across robustness checks":
            "A  Selected Pfam–AEF associations across sensitivity checks",
        "C  Retention by check": "C  Consistency across checks",
    }
    for old, new in substitutions.items():
        if source.count(old) != 1:
            raise RuntimeError(f"Expected one Figure S3 title: {old}")
        source = source.replace(old, new)
    namespace = {"__name__": "corrected_figure_s3_source", "__file__": str(SOURCE_PLOT_SCRIPT)}
    exec(compile(source, str(SOURCE_PLOT_SCRIPT), "exec"), namespace)
    plotter = namespace.get("create_robustness_figure")
    if plotter is None:
        raise RuntimeError("Figure S3 plot function was not loaded")

    non_null = pd.read_csv(SOURCE_AEF_NON_NULL)
    corrected = pd.read_csv(SOURCE_AEF_NULL)
    candidates = pd.read_csv(SOURCE_AEF_CANDIDATES)
    if len(non_null) != 2856 or len(corrected) != 204 or len(candidates) != 68:
        raise RuntimeError(
            f"Unexpected Figure S3 inputs: {len(non_null)}, {len(corrected)}, {len(candidates)}"
        )
    null_for_plot = corrected.copy()
    null_for_plot["empirical_two_sided_p"] = null_for_plot[
        "empirical_two_sided_p_conditional_tail"
    ]
    null_for_plot["selected_AEF_empirical_bh_q"] = null_for_plot[
        "selected_AEF_empirical_bh_q_conditional_tail"
    ]
    plotter(non_null, null_for_plot, FIGURE_S3_PDF, FIGURE_S3_SVG)

    formatted_caption = display_text("Figure S3. " + FIGURE_S3_LEGEND)
    FIGURE_S3_CAPTION.write_text(formatted_caption + "\n", encoding="utf-8")
    tex_builder = load_module(SOURCE_TEX_BUILDER, "figure_s3_tex_builder_v7")
    FIGURE_S3_PAGE_TEX.write_text(
        tex_builder.build_tex(formatted_caption, FIGURE_S3_PDF.name),
        encoding="utf-8",
    )
    run_checked(
        [
            "tectonic",
            "--only-cached",
            "--chatter",
            "minimal",
            "--outdir",
            str(SUPPLEMENT),
            str(FIGURE_S3_PAGE_TEX),
        ],
        cwd=SUPPLEMENT,
    )
    for path in (FIGURE_S3_PDF, FIGURE_S3_PAGE_PDF):
        run_checked(["qpdf", "--check", str(path)])
        if pdf_pages(path) != 1:
            raise RuntimeError(f"Expected a one-page PDF: {path}")

    svg_text = FIGURE_S3_SVG.read_text(encoding="utf-8")
    font_sizes = Counter(re.findall(r"font:\s*([0-9.]+)px", svg_text))
    stroke_widths = Counter(re.findall(r"stroke-width:\s*([0-9.]+)", svg_text))
    if set(font_sizes) != {"6"}:
        raise RuntimeError(f"Figure S3 contains non-6-px SVG text: {font_sizes}")
    if set(stroke_widths) != {"0.25"}:
        raise RuntimeError(f"Figure S3 contains non-0.25 SVG stroke widths: {stroke_widths}")
    pdf_text = " ".join(
        run_checked(["pdftotext", "-layout", str(FIGURE_S3_PAGE_PDF), "-"]).stdout.split()
    ).replace("ﬁ", "fi").replace("ﬂ", "fl")
    for phrase in (
        "99,999 permutations",
        "p/q values were 8.0e-5/0.00272 and 2.0e-5/0.00136",
        "PF10988–A36 and PF13411–A18 each passed five of seven",
    ):
        if phrase not in pdf_text:
            raise RuntimeError(f"Corrected Figure S3 supporting page lacks: {phrase}")
    for stale in (r"(?<!9)9,999 permutations", r"six of seven", r"p=0\.0001", r"q=0\.0034"):
        if re.search(stale, pdf_text):
            raise RuntimeError(f"Corrected Figure S3 supporting page retains stale text: {stale}")
    return {
        "non_null_rows": len(non_null),
        "corrected_null_rows": len(corrected),
        "candidate_rows": len(candidates),
        "permutations": sorted(corrected["permutations"].unique().astype(int).tolist()),
        "all_seven_pairs": sorted(
            (
                candidates.loc[candidates["robust_candidate_all_required_checks"].astype(bool), "pfam"]
                + "--"
                + candidates.loc[candidates["robust_candidate_all_required_checks"].astype(bool), "latent_axis"]
            ).tolist()
        ),
        "svg_font_sizes_px": dict(font_sizes),
        "svg_stroke_widths": dict(stroke_widths),
        "pdf_fonts": run_checked(["pdffonts", str(FIGURE_S3_PDF)]).stdout,
        "supporting_page_text_audit": "PASS",
    }


def copy_figure_s4() -> dict[str, Any]:
    shutil.copy2(SOURCE_FIGURE_S4_PDF, FIGURE_S4_PDF)
    shutil.copy2(SOURCE_FIGURE_S4_SVG, FIGURE_S4_SVG)
    for source, output in (
        (SOURCE_FIGURE_S4_PDF, FIGURE_S4_PDF),
        (SOURCE_FIGURE_S4_SVG, FIGURE_S4_SVG),
    ):
        if sha256(source) != sha256(output):
            raise RuntimeError(f"Figure S4 byte-preserving copy failed: {output}")
    run_checked(["qpdf", "--check", str(FIGURE_S4_PDF)])
    if pdf_pages(FIGURE_S4_PDF) != 1:
        raise RuntimeError("Figure S4 must be one page")
    svg_text = FIGURE_S4_SVG.read_text(encoding="utf-8")
    font_sizes = Counter(re.findall(r"font:\s*([0-9.]+)px", svg_text))
    stroke_widths = Counter(re.findall(r"stroke-width:\s*([0-9.]+)", svg_text))
    if set(font_sizes) != {"6"} or set(stroke_widths) != {"0.25"}:
        raise RuntimeError(
            f"Figure S4 protocol audit failed: font={font_sizes}, stroke={stroke_widths}"
        )
    return {
        "byte_identical_pdf": True,
        "byte_identical_svg": True,
        "svg_font_sizes_px": dict(font_sizes),
        "svg_stroke_widths": dict(stroke_widths),
        "pdf_fonts": run_checked(["pdffonts", str(FIGURE_S4_PDF)]).stdout,
    }


def validate_workbooks() -> dict[str, Any]:
    expected = {
        SOURCE_TABLE_S2: [
            "README",
            "Table S2A Exact-ID GEE",
            "Table S2B Summary",
            "Table S2C FDR pairs",
            "Table S2D File Index",
            "S2E GEE sensitivity",
            "S2F Structured null",
            "S2G Candidate summary",
        ],
        SOURCE_TABLE_S3: [
            "README",
            "S3A_pooled_AEF_Pfam",
            "S3B_phylum_centered",
            "S3C_site_recorded_metadata",
            "S3D_site_axis_summary",
            "S3E_output_index",
            "S3F_within_phylum_summary",
            "S3G_AEF_GEE_alignment",
            "S3H_GEE_variable_summary",
            "S3I_AEF_axis_alignment",
            "S3J_corrected_AEF_null",
            "S3K_AEF_candidate_checks",
        ],
    }
    records: dict[str, Any] = {}
    for path, expected_sheets in expected.items():
        workbook = load_workbook(path, read_only=True, data_only=True)
        if workbook.sheetnames != expected_sheets:
            raise RuntimeError(f"Unexpected worksheet order for {path}: {workbook.sheetnames}")
        records[str(path.resolve())] = {
            "sha256": sha256(path),
            "sheets": {
                sheet.title: {"rows": sheet.max_row, "columns": sheet.max_column}
                for sheet in workbook.worksheets
            },
        }
        workbook.close()
    expected_rows = {
        (SOURCE_TABLE_S2, "S2E GEE sensitivity"): 2017,
        (SOURCE_TABLE_S2, "S2F Structured null"): 85,
        (SOURCE_TABLE_S2, "S2G Candidate summary"): 85,
        (SOURCE_TABLE_S3, "S3G_AEF_GEE_alignment"): 833,
        (SOURCE_TABLE_S3, "S3H_GEE_variable_summary"): 14,
        (SOURCE_TABLE_S3, "S3I_AEF_axis_alignment"): 65,
        (SOURCE_TABLE_S3, "S3J_corrected_AEF_null"): 205,
        (SOURCE_TABLE_S3, "S3K_AEF_candidate_checks"): 69,
    }
    for (path, sheet), rows in expected_rows.items():
        observed = records[str(path.resolve())]["sheets"][sheet]["rows"]
        if observed != rows:
            raise RuntimeError(f"Unexpected row count for {sheet}: {observed} != {rows}")
    return records


def render_legend_pdf() -> dict[str, Any]:
    soffice = next((path for path in SOFFICE_CANDIDATES if path.is_file()), None)
    if soffice is None:
        resolved = shutil.which("soffice")
        if resolved is None:
            raise FileNotFoundError("LibreOffice soffice was not found")
        soffice = Path(resolved)
    with tempfile.TemporaryDirectory(prefix="supp_v7_profile_") as profile, tempfile.TemporaryDirectory(
        prefix="supp_v7_pdf_"
    ) as outdir:
        result = run_checked(
            [
                str(soffice),
                f"-env:UserInstallation={Path(profile).resolve().as_uri()}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                outdir,
                str(OUTPUT_DOCX),
            ]
        )
        produced = Path(outdir) / OUTPUT_PDF.name
        if not produced.is_file():
            raise RuntimeError(f"LibreOffice did not render {produced}")
        shutil.move(produced, OUTPUT_PDF)
    run_checked(["qpdf", "--check", str(OUTPUT_PDF)])
    pages = pdf_pages(OUTPUT_PDF)
    if pages not in {3, 4}:
        raise RuntimeError(f"Unexpected V7 page count: {pages}")
    pdf_text = " ".join(
        run_checked(["pdftotext", "-layout", str(OUTPUT_PDF), "-"]).stdout.split()
    ).replace("ﬁ", "fi").replace("ﬂ", "fl")
    for phrase in (
        MANUSCRIPT_TITLE,
        "Figure S4. Secondary exploratory AEF–Pfam correlation-profile landscape",
        "Table S2G. GEE seven-specification candidate summary",
        "Table S3G. Site-level AEF–GEE alignment",
        "Table S3K. Corrected selected-AEF seven-specification summary",
    ):
        if phrase not in pdf_text:
            raise RuntimeError(f"Rendered V7 PDF lacks: {phrase}")
    return {
        "soffice": str(soffice),
        "conversion_stdout": result.stdout.strip(),
        "conversion_stderr": result.stderr.strip(),
        "pages": pages,
        "pdfinfo": run_checked(["pdfinfo", str(OUTPUT_PDF)]).stdout,
        "pdffonts": run_checked(["pdffonts", str(OUTPUT_PDF)]).stdout,
        "required_text_present": True,
    }


def main() -> None:
    required_inputs = (
        SOURCE_DOCX,
        SOURCE_TXT,
        SOURCE_PDF,
        SOURCE_TABLE_S2,
        SOURCE_TABLE_S3,
        SOURCE_PLOT_SCRIPT,
        SOURCE_AEF_NON_NULL,
        SOURCE_AEF_NULL,
        SOURCE_AEF_CANDIDATES,
        SOURCE_TEX_BUILDER,
        SOURCE_FIGURE_S4_PDF,
        SOURCE_FIGURE_S4_SVG,
    )
    for path in required_inputs:
        if not path.is_file():
            raise FileNotFoundError(path)
    outputs = (
        OUTPUT_DOCX,
        OUTPUT_PDF,
        OUTPUT_TXT,
        OUTPUT_AUDIT,
        FIGURE_S3_PDF,
        FIGURE_S3_SVG,
        FIGURE_S3_CAPTION,
        FIGURE_S3_PAGE_TEX,
        FIGURE_S3_PAGE_PDF,
        FIGURE_S4_PDF,
        FIGURE_S4_SVG,
    )
    existing = [str(path) for path in outputs if path.exists()]
    if existing:
        raise FileExistsError("Refusing to overwrite versioned outputs: " + ", ".join(existing))

    workbook_audit = validate_workbooks()
    docx_audit = build_legend_document()
    figure_s3_audit = build_figure_s3()
    figure_s4_audit = copy_figure_s4()
    render_audit = render_legend_pdf()

    output_files = (
        OUTPUT_DOCX,
        OUTPUT_PDF,
        OUTPUT_TXT,
        FIGURE_S3_PDF,
        FIGURE_S3_SVG,
        FIGURE_S3_CAPTION,
        FIGURE_S3_PAGE_TEX,
        FIGURE_S3_PAGE_PDF,
        FIGURE_S4_PDF,
        FIGURE_S4_SVG,
    )
    audit = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "generator": {
            "path": str(SCRIPT),
            "sha256": sha256(SCRIPT),
            "python": platform.python_version(),
            "pandas": pd.__version__,
            "numpy": np.__version__,
            "matplotlib": matplotlib.__version__,
            "lxml": etree.LXML_VERSION,
        },
        "inputs": {str(path.resolve()): file_record(path) for path in required_inputs},
        "outputs": {str(path.resolve()): file_record(path) for path in output_files},
        "workbooks": workbook_audit,
        "legend_document": docx_audit,
        "figure_s3": figure_s3_audit,
        "figure_s4": figure_s4_audit,
        "rendering": render_audit,
        "scientific_integrity": {
            "synthetic_or_placeholder_data_used": False,
            "scientific_values_hardcoded_as_substitutes_for_analysis": False,
            "completed_results_loaded": True,
            "figure_s3_non_null_values_reused_from_completed_output": True,
            "figure_s3_null_values_loaded_from_corrected_99999_permutation_output": True,
            "figure_s4_vector_copied_byte_for_byte": True,
            "source_v6_preserved": True,
            "source_workbooks_preserved": True,
        },
        "hierarchy": {
            "primary": "named Google Earth Engine variables",
            "secondary": "AlphaEarth Foundations unitless latent axes",
            "axis_interpretation": "site-level external alignment/crosswalk, not semantic decoding",
        },
    }
    OUTPUT_AUDIT.write_text(
        json.dumps(audit, indent=2, sort_keys=True, ensure_ascii=False, default=list) + "\n",
        encoding="utf-8",
    )
    for path in (*output_files, OUTPUT_AUDIT):
        print(path.resolve())


if __name__ == "__main__":
    main()
