#!/usr/bin/env python3
"""Build the final GEE-primary Supplemental Information package (V8).

The builder preserves V7 and uses authenticated completed result tables. It
re-renders Figure S3 so structured-null fill reflects corrected selected-family
BH q < 0.05, distinguishes the seven required gates from additional cells, and
labels Panel C as a nominal-p summary. No scientific values are simulated,
imputed, reconstructed from summaries, or hardcoded as substitutes for data.

Generated: 2026-07-19.
"""

from __future__ import annotations

import hashlib
import importlib.util
import json
import platform
import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib
import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from matplotlib.colors import LinearSegmentedColormap, TwoSlopeNorm
from matplotlib.lines import Line2D
from openpyxl import load_workbook


STAMP = "20260719_214546"
SCRIPT = Path(__file__).resolve()
ROOT = SCRIPT.parents[2]
SUPPLEMENT = SCRIPT.parent
STATS = ROOT / "ISCIENCE_REVISION_20260711" / "analysis_stats"
AEF_RUN = (
    ROOT
    / "ISCIENCE_REVISION_20260711"
    / "aef"
    / "full_aef_corrected_run_20260718_154213"
)

V7_BUILDER = SUPPLEMENT / "build_supplemental_legends_v7_20260719_211018.py"
SOURCE_STEM = SUPPLEMENT / "Supplemental_Information_revised_legends_20260719_211018_V7"
SOURCE_DOCX = SOURCE_STEM.with_suffix(".docx")
SOURCE_TXT = SOURCE_STEM.with_suffix(".txt")
SOURCE_PDF = SOURCE_STEM.with_suffix(".pdf")
SOURCE_TABLE_S2 = (
    SUPPLEMENT / "Table_S2_revised_exact_id_GEE_validation_20260719_215020_V6.xlsx"
)
SOURCE_TABLE_S3 = SUPPLEMENT / "Table_S3_AEF_20260719_214300_V8.xlsx"

SOURCE_AEF_NON_NULL = STATS / "archived_AEF_priority_robustness_20260711_131706.csv"
SOURCE_AEF_NULL = STATS / "corrected_AEF_structured_null_conditional_tail_20260719_210254.csv"
SOURCE_AEF_CANDIDATES = STATS / "corrected_AEF_robust_candidate_table_20260719_210254.csv"
SOURCE_TEX_BUILDER = SUPPLEMENT / "build_figure_s3_supporting_package_20260711_175934.py"
SOURCE_FIGURE_S4_PDF = AEF_RUN / "Figure5_corrected_AEF_Pfam_landscape.pdf"
SOURCE_FIGURE_S4_SVG = AEF_RUN / "Figure5_corrected_AEF_Pfam_landscape.svg"

OUTPUT_STEM = SUPPLEMENT / f"Supplemental_Information_revised_legends_{STAMP}_V8"
OUTPUT_DOCX = OUTPUT_STEM.with_suffix(".docx")
OUTPUT_PDF = OUTPUT_STEM.with_suffix(".pdf")
OUTPUT_TXT = OUTPUT_STEM.with_suffix(".txt")
OUTPUT_AUDIT = SUPPLEMENT / f"Supplemental_Information_revised_legends_{STAMP}_V8_integrity.json"

FIGURE_S3_STEM = f"Figure_S3_posthoc_sensitivity_checks_V4_{STAMP}"
FIGURE_S3_PDF = SUPPLEMENT / f"{FIGURE_S3_STEM}.pdf"
FIGURE_S3_SVG = SUPPLEMENT / f"{FIGURE_S3_STEM}.svg"
FIGURE_S3_CAPTION = SUPPLEMENT / f"Figure_S3_posthoc_sensitivity_checks_caption_V4_{STAMP}.txt"
FIGURE_S3_PAGE_TEX = SUPPLEMENT / f"Figure_S3_posthoc_sensitivity_checks_supporting_page_V5_{STAMP}.tex"
FIGURE_S3_PAGE_PDF = FIGURE_S3_PAGE_TEX.with_suffix(".pdf")

FIGURE_S4_PDF = SUPPLEMENT / f"Figure_S4_exploratory_AEF_Pfam_correlation_profile_landscape_{STAMP}.pdf"
FIGURE_S4_SVG = SUPPLEMENT / f"Figure_S4_exploratory_AEF_Pfam_correlation_profile_landscape_{STAMP}.svg"

SOFFICE_CANDIDATES = (
    Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
    Path("/opt/homebrew/bin/soffice"),
)

FIGURE_S3_LEGEND = (
    "Secondary selected-set sensitivity analysis of protein family (Pfam)–AlphaEarth Foundations (AEF) "
    "latent-feature associations. Panel A displays rank-association effects for 12 selected pairs. The first "
    "seven columns are the required gates: unique-site raw counts, total-hit normalization, and peptide "
    "normalization; the total-hit quality/phylum site-cluster model; the total-hit Benchmarking Universal "
    "Single-Copy Orthologs (BUSCO) ≥50% unique-site analysis; the total-hit three-phylogenetic-eigenvector "
    "topology model; and the total-hit structured null. Subsequent columns show additional sensitivity cells "
    "and genome-level descriptive estimates. Asterisks indicate Benjamini–Hochberg (BH) q < 0.05 across the "
    "68-pair selected family for the indicated method and abundance representation. Panel B shows the "
    "total-hit site-mean effect and 2.5th–97.5th percentiles from 99,999 intact-site permutations within site "
    "phylum-composition strata. Filled red points indicate corrected conditional-tail BH q < 0.05; open blue "
    "points indicate q ≥ 0.05. PF10988–A36 (p = 0.00304, q = 0.05168) and PF00092–A49 (p = 0.01626, "
    "q = 0.221136) are open. Panel C counts the 12 displayed pairs with nominal, unadjusted p < 0.05 at each "
    "displayed cell; these descriptive counts are separate from seven-gate candidate retention. PF01638–A52 "
    "and PF01638–A53 retained direction and selected-family support across all seven required gates (7/7); "
    "PF10988–A36 and PF13411–A18 each passed five of seven (5/7). R denotes raw count, T total-hit "
    "normalization, P peptide normalization, Site unique-site Spearman correlation, Qual the quality/phylum "
    "site-cluster model, B50 and B70 the BUSCO ≥50% and ≥70% unique-site analyses, P3/P5/P10 the three-, "
    "five-, and ten-phylogenetic-eigenvector topology models, Null the structured null, and Geno the "
    "genome-level descriptive Spearman correlation. A00–A63 denote unitless latent axes."
)

FIGURE_S4_LEGEND = (
    "Secondary exploratory AEF–Pfam correlation-profile landscape. Average-linkage hierarchical clustering "
    "with correlation distance organizes Pfam–AEF Spearman profiles across 126 exact-genome-identifier (ID) "
    "records. All 10,707 strict Pfam accessions and 64 axes are displayed as row-wise z scores. The "
    "exact-profile audit identified 8,713 unique profiles and 1,994 additional accessions sharing an exact "
    "profile; all 10,707 accessions remain in the display. The top bar gives the number of global-BH-supported "
    "Pfam pairs per axis, and the right strip gives each Pfam's maximum absolute Spearman coefficient. AEF "
    "axes A00–A63 are unitless latent features. Figure S3 reports the selected-set AEF sensitivity analysis. "
    "Table S3G–I reports a descriptive site-level cross-representation alignment/crosswalk with named Google "
    "Earth Engine variables; axis meanings are distributed and nonunique."
)


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load module: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


V7 = load_module(V7_BUILDER, "supplemental_legends_v7_helpers")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def file_record(path: Path) -> dict[str, Any]:
    return {"path": str(path.resolve()), "bytes": path.stat().st_size, "sha256": sha256(path)}


def run_checked(command: list[str], *, cwd: Path | None = None) -> subprocess.CompletedProcess[str]:
    result = subprocess.run(command, cwd=cwd, capture_output=True, text=True, check=False)
    if result.returncode != 0:
        raise RuntimeError(
            f"Command failed ({result.returncode}): {' '.join(command)}\n"
            f"STDOUT:\n{result.stdout}\nSTDERR:\n{result.stderr}"
        )
    return result


def pdf_pages(path: Path) -> int:
    info = run_checked(["pdfinfo", str(path)]).stdout
    match = re.search(r"^Pages:\s+(\d+)$", info, flags=re.MULTILINE)
    if match is None:
        raise RuntimeError(f"Could not determine PDF page count: {path}")
    return int(match.group(1))


def configure_figure_style() -> None:
    mpl.rcParams.update(
        {
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "svg.fonttype": "none",
            "font.family": "sans-serif",
            "font.sans-serif": ["Arial", "Helvetica"],
            "font.size": 6,
            "axes.labelsize": 6,
            "axes.titlesize": 6,
            "xtick.labelsize": 6,
            "ytick.labelsize": 6,
            "legend.fontsize": 6,
            "axes.linewidth": 0.25,
            "xtick.major.width": 0.25,
            "ytick.major.width": 0.25,
            "xtick.major.size": 2,
            "ytick.major.size": 2,
            "lines.linewidth": 0.25,
            "patch.linewidth": 0.25,
            "grid.linewidth": 0.25,
            "axes.labelpad": 1,
            "xtick.major.pad": 1,
            "ytick.major.pad": 1,
        }
    )


def selected_display_pairs(non_null: pd.DataFrame) -> pd.DataFrame:
    raw = non_null[
        (non_null["representation"] == "raw_count")
        & (non_null["method"] == "site_mean_spearman")
    ].sort_values("p_value", kind="stable")
    explicit = raw[raw["pfam"] != "PF00092"][["pfam", "latent_axis"]].drop_duplicates()
    pf00092 = raw[raw["pfam"] == "PF00092"].head(8)[["pfam", "latent_axis"]]
    selected = pd.concat([explicit, pf00092], ignore_index=True).drop_duplicates().head(12)
    expected = [
        ("PF01638", "A53"),
        ("PF01638", "A52"),
        ("PF10988", "A36"),
        ("PF13411", "A18"),
        ("PF00092", "A49"),
        ("PF00092", "A09"),
        ("PF00092", "A31"),
        ("PF00092", "A21"),
        ("PF00092", "A36"),
        ("PF00092", "A50"),
        ("PF00092", "A35"),
        ("PF00092", "A17"),
    ]
    observed = list(map(tuple, selected.to_numpy()))
    if observed != expected:
        raise RuntimeError(f"Unexpected Figure S3 display set: {observed}")
    return selected


def figure_columns() -> list[dict[str, Any]]:
    required = [
        ("Site\nR", "site_mean_spearman", "raw_count"),
        ("Site\nT", "site_mean_spearman", "per_total_pfam_hit"),
        ("Site\nP", "site_mean_spearman", "per_final_peptide_record"),
        ("Qual\nT", "quality_phylum_sitecluster", "per_total_pfam_hit"),
        ("B50\nT", "site_mean_busco_ge50_spearman", "per_total_pfam_hit"),
        ("P3\nT", "phylo_pev_brokenstick_quality_sitecluster", "per_total_pfam_hit"),
        ("Null\nT", "corrected_structured_null", "per_total_pfam_hit"),
    ]
    additional = [
        ("Qual\nR", "quality_phylum_sitecluster", "raw_count"),
        ("Qual\nP", "quality_phylum_sitecluster", "per_final_peptide_record"),
        ("B50\nR", "site_mean_busco_ge50_spearman", "raw_count"),
        ("B50\nP", "site_mean_busco_ge50_spearman", "per_final_peptide_record"),
        ("B70\nR", "site_mean_busco_ge70_spearman", "raw_count"),
        ("B70\nT", "site_mean_busco_ge70_spearman", "per_total_pfam_hit"),
        ("B70\nP", "site_mean_busco_ge70_spearman", "per_final_peptide_record"),
        ("P3\nR", "phylo_pev_brokenstick_quality_sitecluster", "raw_count"),
        ("P3\nP", "phylo_pev_brokenstick_quality_sitecluster", "per_final_peptide_record"),
        ("P5\nR", "phylo_pev5_quality_sitecluster", "raw_count"),
        ("P5\nT", "phylo_pev5_quality_sitecluster", "per_total_pfam_hit"),
        ("P5\nP", "phylo_pev5_quality_sitecluster", "per_final_peptide_record"),
        ("P10\nR", "phylo_pev10_quality_sitecluster", "raw_count"),
        ("P10\nT", "phylo_pev10_quality_sitecluster", "per_total_pfam_hit"),
        ("P10\nP", "phylo_pev10_quality_sitecluster", "per_final_peptide_record"),
    ]
    descriptive = [
        ("Geno\nR", "genome_level_spearman_descriptive", "raw_count"),
        ("Geno\nT", "genome_level_spearman_descriptive", "per_total_pfam_hit"),
        ("Geno\nP", "genome_level_spearman_descriptive", "per_final_peptide_record"),
    ]
    columns: list[dict[str, Any]] = []
    for group, values in (
        ("required", required),
        ("additional", additional),
        ("descriptive", descriptive),
    ):
        for label, method, representation in values:
            columns.append(
                {
                    "group": group,
                    "label": label,
                    "method": method,
                    "representation": representation,
                }
            )
    if len(columns) != 25 or sum(c["group"] == "required" for c in columns) != 7:
        raise RuntimeError("Figure S3 column definition is incomplete")
    return columns


def build_figure_matrices(
    non_null: pd.DataFrame,
    corrected_null: pd.DataFrame,
    selected: pd.DataFrame,
    columns: list[dict[str, Any]],
) -> tuple[np.ndarray, np.ndarray, np.ndarray]:
    row_order = list(map(tuple, selected.to_numpy()))
    effects = np.full((len(row_order), len(columns)), np.nan)
    qvalues = np.full_like(effects, np.nan)
    pvalues = np.full_like(effects, np.nan)
    for i, (pfam, axis) in enumerate(row_order):
        for j, column in enumerate(columns):
            if column["method"] == "corrected_structured_null":
                hit = corrected_null[
                    (corrected_null["pfam"] == pfam)
                    & (corrected_null["latent_axis"] == axis)
                    & (corrected_null["representation"] == column["representation"])
                ]
                if len(hit) != 1:
                    raise RuntimeError(f"Expected one corrected-null row for {pfam}/{axis}")
                effects[i, j] = hit["observed_spearman_rho"].iloc[0]
                qvalues[i, j] = hit[
                    "selected_AEF_empirical_bh_q_conditional_tail"
                ].iloc[0]
                pvalues[i, j] = hit["empirical_two_sided_p_conditional_tail"].iloc[0]
            else:
                hit = non_null[
                    (non_null["pfam"] == pfam)
                    & (non_null["latent_axis"] == axis)
                    & (non_null["method"] == column["method"])
                    & (non_null["representation"] == column["representation"])
                ]
                if len(hit) != 1:
                    raise RuntimeError(
                        f"Expected one non-null row for {pfam}/{axis}/{column['method']}/"
                        f"{column['representation']}; found {len(hit)}"
                    )
                effects[i, j] = hit["effect"].iloc[0]
                qvalues[i, j] = hit["selected_AEF_set_bh_q"].iloc[0]
                pvalues[i, j] = hit["p_value"].iloc[0]
    return effects, qvalues, pvalues


def build_figure_s3() -> dict[str, Any]:
    non_null = pd.read_csv(SOURCE_AEF_NON_NULL)
    corrected_null = pd.read_csv(SOURCE_AEF_NULL)
    candidates = pd.read_csv(SOURCE_AEF_CANDIDATES)
    if len(non_null) != 2856 or len(corrected_null) != 204 or len(candidates) != 68:
        raise RuntimeError(
            f"Unexpected input rows: {len(non_null)}, {len(corrected_null)}, {len(candidates)}"
        )
    selected = selected_display_pairs(non_null)
    columns = figure_columns()
    effects, qvalues, pvalues = build_figure_matrices(
        non_null, corrected_null, selected, columns
    )
    row_order = list(map(tuple, selected.to_numpy()))

    configure_figure_style()
    colors = [
        (0.0, 0.3, 0.7),
        (0.3, 0.5, 0.9),
        (0.7, 0.8, 0.95),
        (0.95, 0.95, 0.95),
        (0.95, 0.8, 0.7),
        (0.9, 0.4, 0.3),
        (0.7, 0.1, 0.1),
    ]
    cmap = LinearSegmentedColormap.from_list("journal_diverging", colors)
    fig = plt.figure(figsize=(7.0, 5.9))
    grid = fig.add_gridspec(
        2,
        2,
        height_ratios=[3.35, 1.65],
        width_ratios=[1.45, 1.0],
        hspace=0.58,
        wspace=0.34,
    )

    ax_heat = fig.add_subplot(grid[0, :])
    vmax = max(0.4, float(np.nanmax(np.abs(effects))))
    image = ax_heat.imshow(
        effects,
        aspect="auto",
        cmap=cmap,
        norm=TwoSlopeNorm(vmin=-vmax, vcenter=0, vmax=vmax),
    )
    ax_heat.set_xticks(range(len(columns)), [column["label"] for column in columns])
    ax_heat.set_yticks(range(len(row_order)), [f"{pfam}/{axis}" for pfam, axis in row_order])
    ax_heat.set_title(
        "A  Selected Pfam–AEF associations: required gates and additional cells",
        loc="left",
        fontweight="bold",
        pad=14,
    )
    for boundary in (6.5, 21.5):
        ax_heat.axvline(boundary, color="black", linewidth=0.25)
    group_specs = [
        (-0.5, 7, "7 required gates", "#d9e8f5"),
        (6.5, 15, "Additional sensitivities", "#eeeeee"),
        (21.5, 3, "Genome descriptive", "#f4ead7"),
    ]
    for start, width, label, color in group_specs:
        ax_heat.add_patch(
            plt.Rectangle(
                (start, -1.28),
                width,
                0.5,
                facecolor=color,
                edgecolor="black",
                linewidth=0.25,
                clip_on=False,
            )
        )
        ax_heat.text(start + width / 2, -1.03, label, ha="center", va="center")
    ax_heat.set_ylim(len(row_order) - 0.5, -1.32)
    for i in range(effects.shape[0]):
        for j in range(effects.shape[1]):
            if np.isfinite(qvalues[i, j]) and qvalues[i, j] < 0.05:
                ax_heat.text(j, i, "*", ha="center", va="center", color="black")
    colorbar = fig.colorbar(image, ax=ax_heat, fraction=0.017, pad=0.008)
    colorbar.set_label("Rank-association effect")
    colorbar.outline.set_linewidth(0.25)

    ax_null = fig.add_subplot(grid[1, 0])
    selected_keys = set(row_order)
    null_total = corrected_null[
        (corrected_null["representation"] == "per_total_pfam_hit")
        & corrected_null[["pfam", "latent_axis"]].apply(tuple, axis=1).isin(selected_keys)
    ].copy()
    null_total["key"] = list(zip(null_total["pfam"], null_total["latent_axis"]))
    null_total = null_total.set_index("key").reindex(row_order)
    if null_total["observed_spearman_rho"].isna().any():
        raise RuntimeError("Corrected structured-null panel has missing displayed rows")
    ypos = np.arange(len(row_order))
    ax_null.hlines(
        ypos,
        null_total["null_q025"],
        null_total["null_q975"],
        color="#666666",
        linewidth=0.25,
    )
    supported = null_total["selected_AEF_empirical_bh_q_conditional_tail"] < 0.05
    ax_null.scatter(
        null_total.loc[~supported, "observed_spearman_rho"],
        ypos[~supported],
        s=9,
        facecolor="white",
        edgecolor="#1f4e79",
        linewidth=0.25,
        zorder=3,
    )
    ax_null.scatter(
        null_total.loc[supported, "observed_spearman_rho"],
        ypos[supported],
        s=9,
        facecolor="#b22222",
        edgecolor="#b22222",
        linewidth=0.25,
        zorder=3,
    )
    ax_null.axvline(0, color="black", linewidth=0.25)
    ax_null.set_yticks(ypos, [f"{pfam}/{axis}" for pfam, axis in row_order])
    ax_null.invert_yaxis()
    ax_null.set_xlabel("Site-mean rho; corrected null 2.5th–97.5th percentiles")
    ax_null.set_title("B  Corrected structured null", loc="left", fontweight="bold")
    handles = [
        Line2D(
            [0],
            [0],
            marker="o",
            linestyle="none",
            markersize=3,
            markerfacecolor="#b22222",
            markeredgecolor="#b22222",
            markeredgewidth=0.25,
            label="Corrected BH q<0.05",
        ),
        Line2D(
            [0],
            [0],
            marker="o",
            linestyle="none",
            markersize=3,
            markerfacecolor="white",
            markeredgecolor="#1f4e79",
            markeredgewidth=0.25,
            label="q≥0.05",
        ),
    ]
    ax_null.legend(
        handles=handles,
        loc="upper right",
        bbox_to_anchor=(1.0, 1.14),
        ncol=2,
        frameon=False,
        handletextpad=0.3,
        columnspacing=0.8,
        borderaxespad=0,
    )

    ax_count = fig.add_subplot(grid[1, 1])
    nominal_counts = np.sum(np.isfinite(pvalues) & (pvalues < 0.05), axis=0)
    bar_colors = [
        "#4f81bd" if column["group"] == "required" else
        "#999999" if column["group"] == "additional" else "#c49a6c"
        for column in columns
    ]
    ax_count.bar(
        np.arange(len(columns)), nominal_counts, color=bar_colors, width=0.8, linewidth=0
    )
    ax_count.set_xticks(
        np.arange(len(columns)),
        [column["label"].replace("\n", "-") for column in columns],
        rotation=90,
        ha="center",
    )
    ax_count.set_ylabel("Displayed pairs with nominal p<0.05")
    ax_count.set_title("C  Nominal-p summaries", loc="left", fontweight="bold")
    ax_count.set_ylim(0, max(1, int(nominal_counts.max()) + 1))
    ax_count.grid(axis="y", color="#cccccc", linewidth=0.25)
    for boundary in (6.5, 21.5):
        ax_count.axvline(boundary, color="black", linewidth=0.25)

    for axis in (ax_heat, ax_null, ax_count):
        axis.tick_params(width=0.25)
        for spine in axis.spines.values():
            spine.set_linewidth(0.25)
    fig.savefig(
        FIGURE_S3_PDF,
        format="pdf",
        bbox_inches="tight",
        transparent=True,
        edgecolor="none",
    )
    fig.savefig(
        FIGURE_S3_SVG,
        format="svg",
        bbox_inches="tight",
        transparent=True,
        edgecolor="none",
    )
    plt.close(fig)

    caption = "Figure S3. " + FIGURE_S3_LEGEND
    FIGURE_S3_CAPTION.write_text(caption + "\n", encoding="utf-8")
    tex_builder = load_module(SOURCE_TEX_BUILDER, "figure_s3_tex_builder_v8")
    FIGURE_S3_PAGE_TEX.write_text(
        tex_builder.build_tex(caption, FIGURE_S3_PDF.name), encoding="utf-8"
    )
    run_checked(
        [
            "tectonic",
            "--only-cached",
            "--chatter",
            "minimal",
            "--outdir",
            str(SUPPLEMENT),
            str(FIGURE_S3_PAGE_TEX),
        ],
        cwd=SUPPLEMENT,
    )
    for path in (FIGURE_S3_PDF, FIGURE_S3_PAGE_PDF):
        run_checked(["qpdf", "--check", str(path)])
        if pdf_pages(path) != 1:
            raise RuntimeError(f"Expected a one-page PDF: {path}")

    # Assert the audit-critical point encodings directly from the plotted q criterion.
    point_status = {
        f"{pfam}–{axis}": bool(status)
        for (pfam, axis), status in zip(row_order, supported.to_numpy())
    }
    expected_status = {
        "PF01638–A53": True,
        "PF01638–A52": True,
        "PF10988–A36": False,
        "PF13411–A18": True,
        "PF00092–A49": False,
    }
    for pair, expected in expected_status.items():
        if point_status[pair] is not expected:
            raise RuntimeError(f"Panel B status mismatch for {pair}: {point_status[pair]}")

    svg_text = FIGURE_S3_SVG.read_text(encoding="utf-8")
    font_sizes = Counter(re.findall(r"font:\s*([0-9.]+)px", svg_text))
    stroke_widths = Counter(re.findall(r"stroke-width:\s*([0-9.]+)", svg_text))
    if set(font_sizes) != {"6"}:
        raise RuntimeError(f"Figure S3 contains non-6-px text: {font_sizes}")
    if set(stroke_widths) != {"0.25"}:
        raise RuntimeError(f"Figure S3 contains non-0.25 strokes: {stroke_widths}")
    page_text = " ".join(
        run_checked(["pdftotext", "-layout", str(FIGURE_S3_PAGE_PDF), "-"]).stdout.split()
    ).replace("ﬁ", "fi").replace("ﬂ", "fl")
    required_text = (
        "Filled red points indicate corrected conditional-tail BH q < 0.05",
        "PF10988–A36 (p = 0.00304, q = 0.05168)",
        "PF00092–A49 (p = 0.01626, q = 0.221136)",
        "nominal, unadjusted p < 0.05",
        "seven required gates",
    )
    missing = [phrase for phrase in required_text if phrase not in page_text]
    if missing:
        raise RuntimeError(f"Figure S3 supporting page lacks required definitions: {missing}")
    prohibited = re.compile(
        r"\b(?:decoder|decoding|external|internal)\b|independent[- ]validation",
        flags=re.IGNORECASE,
    )
    if prohibited.search(caption):
        raise RuntimeError("Figure S3 caption retains prohibited interpretation wording")
    return {
        "non_null_rows": len(non_null),
        "corrected_null_rows": len(corrected_null),
        "candidate_rows": len(candidates),
        "displayed_pairs": [f"{p}–{a}" for p, a in row_order],
        "columns": columns,
        "required_gate_columns": 7,
        "additional_sensitivity_columns": 15,
        "genome_descriptive_columns": 3,
        "panel_b_fill_rule": "selected_AEF_empirical_bh_q_conditional_tail < 0.05",
        "panel_b_point_status": point_status,
        "panel_c_rule": "count among 12 displayed pairs with unadjusted p < 0.05 per cell",
        "svg_font_sizes_px": dict(font_sizes),
        "svg_stroke_widths": dict(stroke_widths),
        "pdf_fonts": run_checked(["pdffonts", str(FIGURE_S3_PDF)]).stdout,
    }


def build_legend_document() -> dict[str, Any]:
    document = Document(SOURCE_DOCX)
    source_paragraphs = len(document.paragraphs)
    V7.replace_paragraph(V7.find_unique(document, "Figure S3."), "Figure S3.", FIGURE_S3_LEGEND)
    V7.replace_paragraph(V7.find_unique(document, "Figure S4."), "Figure S4.", FIGURE_S4_LEGEND)
    V7.replace_paragraph(
        V7.find_unique(document, "Table S2."),
        "Table S2. Primary named-variable Google Earth Engine (GEE) discovery and sensitivity analyses.",
        "Exact-genome-identifier (ID) matching joined environmental values and Pfam counts for 126 genomes. "
        "The discovery family comprised 87,123 tests, with 6,700–6,703 estimable Pfams per named variable; "
        "84 associations representing 51 Pfams met global Benjamini–Hochberg (BH) q < 0.05, and 13 met "
        "Bonferroni p < 0.05. The primary retained GEE result set comprises 49 Pfam–environment pairs across "
        "30 Pfam accessions that retained direction and selected-family support in all seven required "
        "specifications; all 13 discovery-Bonferroni associations are retained.",
    )
    V7.replace_paragraph(
        V7.find_unique(document, "Table S2D."),
        "Table S2D. Analysis output index.",
        "Paths, byte sizes, row counts, and 256-bit Secure Hash Algorithm (SHA-256) hashes for the exact-ID "
        "extraction, complete discovery family, selected-pair sensitivity analyses, structured null, and "
        "authenticated InterPro annotations.",
    )
    s2g = V7.find_unique(document, "Table S2G.")
    V7.replace_paragraph(
        s2g,
        "Table S2G. GEE seven-specification candidate summary.",
        "All 84 selected discovery pairs, including the seven required specification flags, direction "
        "consistency, retained-result status, and authenticated accession-level InterPro fields for the "
        "49 retained pairs. The retained set comprises 49 pairs across 30 Pfam accessions.",
    )
    V7.insert_after(
        s2g,
        "Table S2H. Retained-Pfam annotation index.",
        "One authenticated InterPro record for each of the 30 Pfam accessions represented by the 49-pair "
        "primary retained GEE result set. Names and entry types are accession-level database labels.",
    )

    V7.replace_paragraph(
        V7.find_unique(document, "Table S3."),
        "Table S3. Secondary AlphaEarth Foundations (AEF) association and site-alignment analyses.",
        "The pooled Pfam–AEF screen evaluated 10,707 strict Pfam accessions against 64 unitless latent axes "
        "across 126 exact-ID genomes (685,248 tests). The descriptive site-level cross-representation "
        "alignment/crosswalk evaluated 64 axes against 13 named GEE variables at up to 90 coordinate sites "
        "(832 tests); 224 met global BH q < 0.05 and 44 met Bonferroni p < 0.05. Axis meanings are distributed "
        "and nonunique, and Figure S4 displays the complete exploratory Pfam–AEF correlation-profile landscape.",
    )
    V7.replace_paragraph(
        V7.find_unique(document, "Table S3E."),
        "Table S3E. AEF output index.",
        "Paths, byte sizes, row counts, and SHA-256 hashes for the pooled AEF analyses, descriptive site-level "
        "cross-representation AEF–GEE alignment/crosswalk, and corrected selected-set sensitivity outputs.",
    )
    V7.replace_paragraph(
        V7.find_unique(document, "Table S3G."),
        "Table S3G. Site-level AEF–GEE alignment.",
        "All 832 descriptive cross-representation Spearman correlations between 64 unitless AEF axes and "
        "13 named GEE variables after exact-ID joining and unique-site aggregation. Global BH correction spans "
        "all 832 correlations; 224 met q < 0.05 and 44 met Bonferroni p < 0.05. This site-level alignment/"
        "crosswalk retains distributed, nonunique axis meanings.",
    )
    V7.replace_paragraph(
        V7.find_unique(document, "Table S3H."),
        "Table S3H. GEE-variable alignment summary.",
        "For each of the 13 named GEE variables, the table reports the number of coordinate sites, the "
        "strongest AEF-axis correlation, and counts meeting global BH and Bonferroni thresholds in the "
        "832-test descriptive site-level cross-representation alignment/crosswalk.",
    )
    V7.replace_paragraph(
        V7.find_unique(document, "Table S3I."),
        "Table S3I. AEF-axis alignment summary.",
        "For each of A00–A63, the table reports its strongest named GEE correlate and counts meeting global BH "
        "and Bonferroni thresholds in the descriptive site-level cross-representation alignment/crosswalk. "
        "Axis meanings are distributed and nonunique. A52 aligned negatively with all five sea-surface-"
        "temperature (SST) summaries (strongest rho = -0.704 for maximum SST), whereas A53 aligned positively "
        "with all five (rho = 0.604 for mean SST).",
    )

    output_text = V7.document_text(document)
    prohibited = re.compile(
        r"\b(?:decoder|decoding|external|internal)\b|independent[- ]validation",
        flags=re.IGNORECASE,
    )
    found = sorted(set(match.group(0) for match in prohibited.finditer(output_text)))
    if found:
        raise RuntimeError(f"Prohibited interpretation wording remains in V8: {found}")
    required = (
        "49 Pfam–environment pairs across 30 Pfam accessions",
        "Table S2H.",
        "descriptive site-level cross-representation alignment/crosswalk",
        "Axis meanings are distributed and nonunique",
        "Filled red points indicate corrected conditional-tail BH q < 0.05",
        "PF10988–A36 (p = 0.00304, q = 0.05168)",
        "nominal, unadjusted p < 0.05",
    )
    missing = [phrase for phrase in required if phrase not in output_text]
    if missing:
        raise RuntimeError(f"Required V8 wording is absent: {missing}")

    document.save(OUTPUT_DOCX)
    OUTPUT_TXT.write_text(output_text, encoding="utf-8")
    reopened = Document(OUTPUT_DOCX)
    if V7.document_text(reopened) != output_text:
        raise RuntimeError("V8 DOCX visible paragraphs and TXT differ")
    with zipfile.ZipFile(OUTPUT_DOCX) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("V8 DOCX ZIP integrity failed")
        document_xml = archive.read("word/document.xml")
    if b"<w:ins" in document_xml or b"<w:del" in document_xml:
        raise RuntimeError("Tracked-change markup is present in V8")
    return {
        "source_paragraphs": source_paragraphs,
        "output_paragraphs": len(reopened.paragraphs),
        "inserted_paragraphs": len(reopened.paragraphs) - source_paragraphs,
        "docx_txt_visible_text_identity": True,
        "tracked_change_markup_absent": True,
    }


def validate_workbooks() -> dict[str, Any]:
    expected_sheets = {
        SOURCE_TABLE_S2: [
            "README",
            "Table S2A Exact-ID GEE",
            "Table S2B Summary",
            "Table S2C FDR pairs",
            "Table S2D File Index",
            "S2E GEE sensitivity",
            "S2F Structured null",
            "S2G Candidate summary",
            "S2H Retained annotations",
        ],
        SOURCE_TABLE_S3: [
            "README",
            "S3A_pooled_AEF_Pfam",
            "S3B_phylum_centered",
            "S3C_site_recorded_metadata",
            "S3D_site_axis_summary",
            "S3E_output_index",
            "S3F_within_phylum_summary",
            "S3G_AEF_GEE_alignment",
            "S3H_GEE_variable_summary",
            "S3I_AEF_axis_alignment",
            "S3J_corrected_AEF_null",
            "S3K_AEF_candidate_checks",
        ],
    }
    expected_data_rows = {
        (SOURCE_TABLE_S2, "S2E GEE sensitivity"): 2016,
        (SOURCE_TABLE_S2, "S2F Structured null"): 84,
        (SOURCE_TABLE_S2, "S2G Candidate summary"): 84,
        (SOURCE_TABLE_S2, "S2H Retained annotations"): 30,
        (SOURCE_TABLE_S3, "S3G_AEF_GEE_alignment"): 832,
        (SOURCE_TABLE_S3, "S3J_corrected_AEF_null"): 204,
        (SOURCE_TABLE_S3, "S3K_AEF_candidate_checks"): 68,
    }
    records: dict[str, Any] = {}
    for path, sheets in expected_sheets.items():
        workbook = load_workbook(path, read_only=True, data_only=True)
        if workbook.sheetnames != sheets:
            raise RuntimeError(f"Unexpected sheet order for {path}: {workbook.sheetnames}")
        records[str(path.resolve())] = {
            "sha256": sha256(path),
            "sheets": {
                sheet.title: {"data_rows": sheet.max_row - 1, "columns": sheet.max_column}
                for sheet in workbook.worksheets
            },
        }
        workbook.close()
    for (path, sheet), expected in expected_data_rows.items():
        observed = records[str(path.resolve())]["sheets"][sheet]["data_rows"]
        if observed != expected:
            raise RuntimeError(f"Unexpected {sheet} rows: {observed} != {expected}")
    return records


def copy_figure_s4() -> dict[str, Any]:
    shutil.copy2(SOURCE_FIGURE_S4_PDF, FIGURE_S4_PDF)
    shutil.copy2(SOURCE_FIGURE_S4_SVG, FIGURE_S4_SVG)
    for source, output in (
        (SOURCE_FIGURE_S4_PDF, FIGURE_S4_PDF),
        (SOURCE_FIGURE_S4_SVG, FIGURE_S4_SVG),
    ):
        if sha256(source) != sha256(output):
            raise RuntimeError(f"Figure S4 copy differs: {output}")
    run_checked(["qpdf", "--check", str(FIGURE_S4_PDF)])
    return {
        "pdf_byte_identical": True,
        "svg_byte_identical": True,
        "pdf_fonts": run_checked(["pdffonts", str(FIGURE_S4_PDF)]).stdout,
    }


def render_legend_pdf() -> dict[str, Any]:
    soffice = next((path for path in SOFFICE_CANDIDATES if path.is_file()), None)
    if soffice is None:
        resolved = shutil.which("soffice")
        if resolved is None:
            raise FileNotFoundError("LibreOffice soffice was not found")
        soffice = Path(resolved)
    with tempfile.TemporaryDirectory(prefix="supp_v8_profile_") as profile, tempfile.TemporaryDirectory(
        prefix="supp_v8_pdf_"
    ) as outdir:
        result = run_checked(
            [
                str(soffice),
                f"-env:UserInstallation={Path(profile).resolve().as_uri()}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                outdir,
                str(OUTPUT_DOCX),
            ]
        )
        produced = Path(outdir) / OUTPUT_PDF.name
        if not produced.is_file():
            raise RuntimeError(f"LibreOffice did not create {produced}")
        shutil.move(produced, OUTPUT_PDF)
    run_checked(["qpdf", "--check", str(OUTPUT_PDF)])
    pages = pdf_pages(OUTPUT_PDF)
    if pages not in {3, 4}:
        raise RuntimeError(f"Unexpected V8 PDF pages: {pages}")
    text = " ".join(
        run_checked(["pdftotext", "-layout", str(OUTPUT_PDF), "-"]).stdout.split()
    ).replace("ﬁ", "fi").replace("ﬂ", "fl")
    for phrase in (
        "49 Pfam–environment pairs across 30 Pfam accessions",
        "Table S2H. Retained-Pfam annotation index",
        "Filled red points indicate corrected conditional-tail BH q < 0.05",
        "site-level cross-representation alignment/crosswalk",
    ):
        if phrase not in text:
            raise RuntimeError(f"Rendered V8 PDF lacks: {phrase}")
    return {
        "soffice": str(soffice),
        "stdout": result.stdout.strip(),
        "stderr": result.stderr.strip(),
        "pages": pages,
        "pdfinfo": run_checked(["pdfinfo", str(OUTPUT_PDF)]).stdout,
        "pdffonts": run_checked(["pdffonts", str(OUTPUT_PDF)]).stdout,
    }


def main() -> None:
    required_inputs = (
        V7_BUILDER,
        SOURCE_DOCX,
        SOURCE_TXT,
        SOURCE_PDF,
        SOURCE_TABLE_S2,
        SOURCE_TABLE_S3,
        SOURCE_AEF_NON_NULL,
        SOURCE_AEF_NULL,
        SOURCE_AEF_CANDIDATES,
        SOURCE_TEX_BUILDER,
        SOURCE_FIGURE_S4_PDF,
        SOURCE_FIGURE_S4_SVG,
    )
    for path in required_inputs:
        if not path.is_file():
            raise FileNotFoundError(path)
    outputs = (
        OUTPUT_DOCX,
        OUTPUT_PDF,
        OUTPUT_TXT,
        OUTPUT_AUDIT,
        FIGURE_S3_PDF,
        FIGURE_S3_SVG,
        FIGURE_S3_CAPTION,
        FIGURE_S3_PAGE_TEX,
        FIGURE_S3_PAGE_PDF,
        FIGURE_S4_PDF,
        FIGURE_S4_SVG,
    )
    existing = [str(path) for path in outputs if path.exists()]
    if existing:
        raise FileExistsError("Refusing to overwrite versioned outputs: " + ", ".join(existing))

    workbook_audit = validate_workbooks()
    document_audit = build_legend_document()
    figure_s3_audit = build_figure_s3()
    figure_s4_audit = copy_figure_s4()
    rendering_audit = render_legend_pdf()

    output_files = (
        OUTPUT_DOCX,
        OUTPUT_PDF,
        OUTPUT_TXT,
        FIGURE_S3_PDF,
        FIGURE_S3_SVG,
        FIGURE_S3_CAPTION,
        FIGURE_S3_PAGE_TEX,
        FIGURE_S3_PAGE_PDF,
        FIGURE_S4_PDF,
        FIGURE_S4_SVG,
    )
    audit = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "generator": {
            "path": str(SCRIPT.resolve()),
            "sha256": sha256(SCRIPT),
            "python": platform.python_version(),
            "pandas": pd.__version__,
            "numpy": np.__version__,
            "matplotlib": matplotlib.__version__,
        },
        "inputs": {str(path.resolve()): file_record(path) for path in required_inputs},
        "outputs": {str(path.resolve()): file_record(path) for path in output_files},
        "workbooks": workbook_audit,
        "document": document_audit,
        "figure_s3": figure_s3_audit,
        "figure_s4": figure_s4_audit,
        "rendering": rendering_audit,
        "scientific_integrity": {
            "synthetic_or_placeholder_data_used": False,
            "scientific_values_hardcoded_as_substitutes_for_analysis": False,
            "completed_result_rows_loaded": True,
            "panel_b_fill_uses_corrected_selected_family_q": True,
            "source_v7_preserved": True,
            "source_workbooks_preserved": True,
            "figure_s4_copied_byte_for_byte": True,
        },
        "hierarchy": {
            "primary": "49 retained named-GEE pairs across 30 Pfam accessions",
            "secondary": "AEF latent-axis analyses",
            "axis_interpretation": (
                "descriptive site-level cross-representation alignment/crosswalk; "
                "axis meanings are distributed and nonunique"
            ),
        },
    }
    OUTPUT_AUDIT.write_text(
        json.dumps(audit, indent=2, sort_keys=True, ensure_ascii=False, default=list) + "\n",
        encoding="utf-8",
    )
    for path in (*output_files, OUTPUT_AUDIT):
        print(f"{path.resolve()}\t{sha256(path)}")


if __name__ == "__main__":
    main()
